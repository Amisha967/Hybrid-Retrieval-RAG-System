import os
import io
from typing import Tuple, Dict, Any

def extract_text_from_file(file_bytes: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """
    Extracts text and metadata from various file formats (.txt, .md, .pdf, .docx, .json, .csv).
    """
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    metadata: Dict[str, Any] = {
        "filename": filename,
        "file_extension": ext,
        "file_size_bytes": len(file_bytes)
    }

    if ext in [".txt", ".md", ".json", ".csv", ".log", ".yaml", ".yml", ".py"]:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="replace")

    elif ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            metadata["page_count"] = len(reader.pages)
            extracted_pages = []
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    extracted_pages.append(page_text)
            text = "\n\n".join(extracted_pages)
        except Exception as e:
            # Fallback if pypdf fails or is missing
            text = file_bytes.decode("utf-8", errors="replace")
            metadata["parse_warning"] = str(e)

    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            metadata["paragraph_count"] = len(doc.paragraphs)
        except Exception as e:
            text = file_bytes.decode("utf-8", errors="replace")
            metadata["parse_warning"] = str(e)

    else:
        # Generic text attempt
        text = file_bytes.decode("utf-8", errors="replace")

    # Clean text whitespace
    text = text.strip()
    metadata["character_count"] = len(text)
    metadata["word_count"] = len(text.split())

    return text, metadata
